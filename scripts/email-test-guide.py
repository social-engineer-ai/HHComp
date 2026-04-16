"""Generate a Word document guiding the team on how to test the website, email it."""
import base64, json, mimetypes, sys, urllib.parse, urllib.request
from email.message import EmailMessage
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_env(path):
    env = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, v = line.split("=", 1)
        env[k.strip()] = v.strip().strip('"').strip("'")
    return env

ENV = load_env(Path(__file__).resolve().parent.parent / ".env")
RECIPIENT = sys.argv[1] if len(sys.argv) > 1 else "ashishk@illinois.edu"

def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def p(doc, text):
    doc.add_paragraph(text)

def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def step(doc, n, text):
    para = doc.add_paragraph(style="List Number")
    para.add_run(text)

def bold_lead(doc, lead, rest):
    para = doc.add_paragraph()
    r = para.add_run(lead)
    r.bold = True
    para.add_run(rest)


def build(out_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_heading("How to test the competition platform", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A walkthrough for the Gies SCM Analytics Competition 2026 team")
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # Intro
    bold_lead(doc, "Live URL: ", "https://16-59-203-133.sslip.io")
    bold_lead(doc, "Admin login: ", "uiucbadm576@gmail.com / hhcomp-admin-2026")
    bold_lead(doc, "Goal of this testing pass: ", "Exercise every flow a student will hit before April 17 launch, and every admin action organizers will use between April 17 and May 7. Report any failures back to the platform owner.")
    doc.add_paragraph()

    # 1. What you will need
    h1(doc, "1. What you will need")
    bullet(doc, "Two email addresses you control for the two student roles. Plus-addressing works on both @gmail.com and @illinois.edu: use netid+alpha@illinois.edu and netid+beta@illinois.edu as two virtual students, and both messages land in your main netid inbox. This is the easiest way for one person to test the full team flow end to end.")
    bullet(doc, "Heads-up: UIUC spam filtering sometimes flags verification emails from the platform on first encounter. Check your spam or junk folder if you don't see a code within a minute. Once you mark a message as 'Not junk', future messages land in the inbox directly.")
    bullet(doc, "A modern browser (Chrome, Firefox, Safari, or Edge).")
    bullet(doc, "A second browser window or an incognito session so you can hold two separate logins at once.")
    bullet(doc, "A small Excel file (.xlsx) to use as a prediction test upload. Any .xlsx will do; the structural validation only checks for a Part column and numeric columns.")
    bullet(doc, "A small PDF, a small PPTX, and a small code file (or a ZIP) to test the other three submission components.")
    bullet(doc, "The admin login credentials above.")

    # 2. Before you start
    h1(doc, "2. Before you start")
    bullet(doc, "Clear cookies or use an incognito window for each student you are impersonating. Otherwise the browser will keep you logged in as whoever last authenticated.")
    bullet(doc, "Always check spam or junk if a verification or invitation email does not arrive within a minute. As a fallback, the platform logs every outgoing email to the server, so the platform owner can always pull the code manually from the container logs.")
    bullet(doc, "Test data (registrations, submissions, NDA signatures) will be visible to the admin in /admin. After testing, ask the platform owner to run a cleanup.")

    # 3. Student happy path
    h1(doc, "3. Student happy path (run this first)")
    p(doc, "This is the golden path a competition team will follow. Go through all steps in order with one browser session. Each step should work without errors.")

    h2(doc, "3.1 Landing page")
    step(doc, 1, "Open https://16-59-203-133.sslip.io/ in your browser.")
    step(doc, 2, "Confirm you see the red hero section with the LOSI 22S Sprint Car photo, the $1,100 prize callout, the four-step structure, the timeline, and the Register button.")
    step(doc, 3, "Click Register.")

    h2(doc, "3.2 Team Lead registration")
    step(doc, 1, "Fill in your full name, an @illinois.edu email, a password of at least 8 characters, and a unique team name.")
    step(doc, 2, "Tick the Gies eligibility checkbox. Submit.")
    step(doc, 3, "You should be redirected to the verify-email page with your email prefilled.")
    step(doc, 4, "Open your email inbox. You should receive a message from Gies SCM Analytics Competition with a 6-digit verification code. Check your spam folder if it does not arrive within a minute.")
    step(doc, 5, "Enter the code on the verify screen and submit. You should land on the dashboard.")

    h2(doc, "3.3 Invite a teammate")
    step(doc, 1, "On the dashboard, the team card should say Incomplete (1 of 2).")
    step(doc, 2, "In the invite form, enter a second @illinois.edu email (the one you will use for the teammate).")
    step(doc, 3, "Click Send invite. The invitation code appears in a grey box below the form.")
    step(doc, 4, "Open the second email inbox (or a private window). You should receive a team invitation email with the code and a join link.")

    h2(doc, "3.4 Team Member joins")
    step(doc, 1, "In a second browser (or incognito window), open the join link from the invitation email, or navigate to https://16-59-203-133.sslip.io/join.")
    step(doc, 2, "Enter the invitation code, your name, the invited email (must match exactly), a password, and tick the eligibility checkbox.")
    step(doc, 3, "Submit. You should land on the verify-email page. Enter the 6-digit code from your inbox.")
    step(doc, 4, "Both browsers should now show the team as Complete (2 of 2) when you refresh the dashboard.")

    h2(doc, "3.5 NDA signing")
    step(doc, 1, "On each browser (lead and member), click the Review and sign NDA button on the dashboard, or visit /nda.")
    step(doc, 2, "The NDA PDF should render in an embedded viewer on the page.")
    step(doc, 3, "Type your full legal name exactly as you registered it (case-insensitive, whitespace-tolerant).")
    step(doc, 4, "Tick the I agree checkbox and click Sign NDA.")
    step(doc, 5, "You should be redirected back to the dashboard. The NDA status card should show 1 of 2 signed (or 2 of 2 after the second signature).")
    step(doc, 6, "Once both members have signed, both dashboards should update to show NDA complete.")

    h2(doc, "3.6 Data downloads (after admin enables)")
    p(doc, "Data downloads are locked by default. An admin must enable them and upload at least one data file for this step to work. Skip to section 4 (admin tests) if needed, then come back.")
    step(doc, 1, "Visit the dashboard. The Data card should now say Available (assuming the admin has uploaded files and toggled downloads on).")
    step(doc, 2, "Click Go to data downloads, or visit /data.")
    step(doc, 3, "You should see a list of active files with Download buttons. Click one, the file should download.")

    h2(doc, "3.7 Submission (all four components)")
    p(doc, "Submissions can be uploaded before the primary deadline. Each of the four components has its own upload zone. You can replace any file any number of times; the latest version is what gets scored and reviewed.")
    step(doc, 1, "Navigate to /submit.")
    step(doc, 2, "Upload a small .xlsx file in the Prediction file zone. It should appear with a green UPLOADED badge and version number.")
    step(doc, 3, "Upload a .py or .ipynb or .zip file in the Code zone.")
    step(doc, 4, "Upload a .pdf file in the Methodology zone.")
    step(doc, 5, "Upload a .pptx file in the Presentation zone.")
    step(doc, 6, "After all four are uploaded, a green All four components uploaded banner should appear at the top.")
    step(doc, 7, "Replace the prediction file with a different .xlsx to verify the version history feature. Click the version history disclosure to see the prior version listed.")
    step(doc, 8, "Both team members should receive email confirmations for each upload.")

    h2(doc, "3.8 Leaderboard")
    step(doc, 1, "Visit /leaderboard.")
    step(doc, 2, "Once the grader has processed your prediction file (via the cron job the admin runs), you should see your team listed with a weighted MAPE score. Note: until the answer key and grading script are uploaded by the admin, scoring will queue but not complete.")

    # 4. Admin tests
    h1(doc, "4. Admin tests")
    p(doc, "Log in at /login with the admin credentials. You should be redirected to /admin. The sidebar shows all admin sections.")

    h2(doc, "4.1 Dashboard and user management")
    step(doc, 1, "Click Dashboard. Confirm the stats reflect the teams and users you have created during student testing.")
    step(doc, 2, "Click Users & Teams. Confirm all test teams are listed. Try creating a staff account using the form at the top (any email domain). The new user should appear in the staff accounts table.")
    step(doc, 3, "Log out of admin, open /login in a new window, and log in as the newly created manager account to confirm it works. You should land on /admin with limited permissions. Log out when done.")

    h2(doc, "4.2 Content management")
    step(doc, 1, "Log back in as admin. Click Data & Files.")
    step(doc, 2, "Under the NDA subsection, confirm the current NDA PDF is listed with an ACTIVE badge. Try uploading a replacement PDF. The old version should become inactive and the new one active.")
    step(doc, 3, "Under Data files, upload a test dataset (any .xlsx or .csv or .zip). It should appear in the list.")
    step(doc, 4, "Click Enable downloads to toggle the master switch. Confirm that a student browser (in another window) can now see the file at /data after both members have signed the NDA.")
    step(doc, 5, "Under Grader, upload an answer key file. You can skip the grading script for now: the platform ships a built-in weighted MAPE grader that runs when no custom script is uploaded.")

    h2(doc, "4.3 Announcements, FAQ, and broadcast")
    step(doc, 1, "Click Announcements. Create a test announcement with a title and Markdown body. Visit /announcements in another window to confirm it appears for students and on the public landing.")
    step(doc, 2, "Click FAQ. Add a test Q&A pair and confirm it appears at /faq.")
    step(doc, 3, "Click Broadcast. Compose a test message, set scope to All complete teams, tick Also archive as announcement, and send. Confirm both team members receive the email and that the message also appears at /announcements.")
    step(doc, 4, "Click Audit log. Confirm every admin action you have taken is recorded with timestamp, action name, and entity.")

    h2(doc, "4.4 Submissions and leaderboard controls")
    step(doc, 1, "Click Submissions. Confirm your test team is listed with its four uploaded components. Try downloading each file.")
    step(doc, 2, "Click the Finalist toggle on your test team. That team should now show a FINALIST badge on the leaderboard.")
    step(doc, 3, "Click Leaderboard. Try the visibility toggle: set it to FROZEN and confirm students see a frozen notice (reload a student window). Set it back to VISIBLE.")
    step(doc, 4, "Try the score override: enter a new value and reason, click Save. The score should update and appear with an (adjusted) label.")
    step(doc, 5, "Click Export CSV and confirm a CSV file downloads.")

    h2(doc, "4.5 Settings")
    step(doc, 1, "Click Settings. Confirm the three dates (submission deadline, grace end, registration close) are editable. Change one briefly, save, then revert.")

    # 5. Edge cases
    h1(doc, "5. Edge cases to try")
    bullet(doc, "Register with a non @illinois.edu email (should be rejected at the form).")
    bullet(doc, "Try to register with the same email twice (should show a clear error message).")
    bullet(doc, "Try to register a team with a name that already exists (should show a clear error).")
    bullet(doc, "Enter a wrong verification code twice, then enter the correct one (should succeed, counter resets).")
    bullet(doc, "Invite a teammate, then try to invite a different teammate while the first invitation is still pending. The first invite should be revoked.")
    bullet(doc, "On the NDA page, type a name that does not match your registered name. The signing should fail with a clear error.")
    bullet(doc, "Upload a file of the wrong type to a submission zone (e.g. a .pdf to the Prediction slot). Should be rejected.")
    bullet(doc, "Upload the same file twice to the same submission zone. Both versions should appear in version history.")
    bullet(doc, "Deactivate a student account from the admin users page and try to log in as them. Should be rejected.")
    bullet(doc, "Send a broadcast to Finalists only while no team is marked finalist. Should show an error instead of sending to zero people.")

    # 6. Reporting bugs
    h1(doc, "6. Reporting bugs and cleanup")
    p(doc, "When you find something that does not work:")
    bullet(doc, "Note the URL, what you clicked, what you expected to happen, and what actually happened.")
    bullet(doc, "Screenshots are gold. Paste them into your bug report.")
    bullet(doc, "If the browser shows a server error page, capture the full error text if visible.")
    bullet(doc, "Send the list to the platform owner.")
    p(doc, "")
    p(doc, "After the testing pass, the platform owner can run a one-off cleanup to delete all test accounts, teams, submissions, and audit log entries so the platform is fresh for the real launch on April 17.")

    # 7. Quick smoke test
    h1(doc, "7. Five-minute smoke test")
    p(doc, "If you only have a few minutes to confirm the site is up, run this:")
    step(doc, 1, "Visit https://16-59-203-133.sslip.io/. Confirm the landing page loads with the car photo and red hero.")
    step(doc, 2, "Click Register. The form loads without errors.")
    step(doc, 3, "Visit /faq. At least one FAQ item is visible.")
    step(doc, 4, "Visit /login and log in as admin. You should land on /admin with the sidebar.")
    step(doc, 5, "Click Dashboard in the sidebar. Live counts should render.")
    p(doc, "If all five pass, the platform is up and responding.")

    doc.save(str(out_path))
    print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")


def get_access_token():
    data = urllib.parse.urlencode({
        "client_id": ENV["GOOGLE_CLIENT_ID"],
        "client_secret": ENV["GOOGLE_CLIENT_SECRET"],
        "refresh_token": ENV["GMAIL_REFRESH_TOKEN"],
        "grant_type": "refresh_token",
    }).encode()
    req = urllib.request.Request(
        "https://oauth2.googleapis.com/token", data=data,
        headers={"Content-Type": "application/x-www-form-urlencoded"})
    return json.loads(urllib.request.urlopen(req).read())["access_token"]


def send_email(token, to, subject, body, attachment):
    msg = EmailMessage()
    msg["From"] = f"{ENV.get('GMAIL_SENDER_NAME', 'SCM Platform')} <{ENV['GMAIL_SENDER_ADDRESS']}>"
    msg["To"] = to
    msg["Subject"] = subject
    msg.set_content(body)
    with open(attachment, "rb") as fh:
        msg.add_attachment(fh.read(), maintype="application",
                           subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                           filename=attachment.name)
    raw = base64.urlsafe_b64encode(bytes(msg)).decode("ascii")
    req = urllib.request.Request(
        "https://gmail.googleapis.com/gmail/v1/users/me/messages/send",
        data=json.dumps({"raw": raw}).encode(),
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"})
    return json.loads(urllib.request.urlopen(req).read())


out_dir = Path(__file__).resolve().parent.parent / "tmp"
out_dir.mkdir(parents=True, exist_ok=True)
out_path = out_dir / "SCM-Analytics-Competition-2026-Testing-Guide.docx"
build(out_path)

token = get_access_token()
res = send_email(
    token,
    to=RECIPIENT,
    subject="Testing guide: SCM Analytics Competition 2026 platform",
    body=(
        "Hi,\n\n"
        "Attached is a Word document with a step-by-step testing guide for the "
        "competition platform. Share with the team members who should help test "
        "before the April 17 launch.\n\n"
        "Live URL: https://16-59-203-133.sslip.io\n"
        "Admin login: uiucbadm576@gmail.com / hhcomp-admin-2026\n\n"
        "The document covers the full student happy path, admin tests, edge cases, "
        "bug reporting, and a five-minute smoke test.\n"
    ),
    attachment=out_path,
)
print(f"Sent: id={res.get('id')} to={RECIPIENT}")
