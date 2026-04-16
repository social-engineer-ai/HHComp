"""
Generate a Word document describing the platform's flow and features,
and email it via Gmail (using the existing OAuth refresh token) to a
recipient address.

Usage: python scripts/email-features-doc.py [recipient_email]
"""

import base64
import json
import mimetypes
import os
import sys
import urllib.parse
import urllib.request
from email.message import EmailMessage
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---- config (loaded from .env) ----

def load_env(path: Path) -> dict:
    env = {}
    if not path.exists():
        return env
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if "=" not in line:
            continue
        k, v = line.split("=", 1)
        v = v.strip().strip('"').strip("'")
        env[k.strip()] = v
    return env


ENV = load_env(Path(__file__).resolve().parent.parent / ".env")
CLIENT_ID = ENV.get("GOOGLE_CLIENT_ID")
CLIENT_SECRET = ENV.get("GOOGLE_CLIENT_SECRET")
REFRESH_TOKEN = ENV.get("GMAIL_REFRESH_TOKEN")
SENDER = ENV.get("GMAIL_SENDER_ADDRESS", "uiucbadm576@gmail.com")
SENDER_NAME = ENV.get("GMAIL_SENDER_NAME", "Gies Supply Chain Analytics Competition")

RECIPIENT = sys.argv[1] if len(sys.argv) > 1 else "ashishk@illinois.edu"


# ---- build the .docx ----

def build_doc(out_path: Path) -> None:
    doc = Document()

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title
    t = doc.add_heading("Supply Chain Analytics Competition 2026", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Platform flow and feature overview")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # Meta box
    meta = doc.add_paragraph()
    meta.add_run("Competition partner: ").bold = True
    meta.add_run("Horizon Hobby\n")
    meta.add_run("Host: ").bold = True
    meta.add_run("Gies College of Business, University of Illinois Urbana-Champaign\n")
    meta.add_run("Competition window: ").bold = True
    meta.add_run("April 17 to May 7, 2026\n")
    meta.add_run("Live URL: ").bold = True
    meta.add_run("https://16-59-203-133.sslip.io\n")
    meta.add_run("Hosting: ").bold = True
    meta.add_run("Single AWS EC2 instance (us-east-2), Dockerized stack\n")
    meta.add_run("Stack: ").bold = True
    meta.add_run(
        "Next.js 16 + TypeScript + Tailwind (frontend and backend), PostgreSQL, "
        "Caddy reverse proxy with automatic HTTPS via Let's Encrypt, Python FastAPI "
        "grader service, S3 for file storage, Gmail OAuth for transactional email"
    )

    doc.add_paragraph()

    # --- Section 1: What the platform does ---
    doc.add_heading("1. What the platform does", level=1)
    doc.add_paragraph(
        "The platform hosts the Gies College of Business Supply Chain Analytics "
        "Competition, presented by Horizon Hobby. It handles the full lifecycle of a "
        "two-week student forecasting competition:"
    )
    bullets = [
        "Team registration and invitation flow",
        "NDA execution with cryptographic version pinning",
        "Gated distribution of Horizon Hobby's confidential dataset",
        "Four-component deliverable submissions with version history",
        "Automated weighted-MAPE scoring via a sandboxed Python grader",
        "Live leaderboard with freeze, hide, and override controls",
        "Finalist selection and presentation workflow",
        "Broadcast email and announcements for organizer communication",
        "Full audit trail of every administrative action",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # --- Section 2: Key dates ---
    doc.add_heading("2. Competition timeline", level=1)
    timeline = [
        ("April 17 (Friday)", "Competition announced, website live, registration opens"),
        ("April 20 (Monday)", "Kickoff briefing, NDA made available, data package distributed"),
        ("April 27 (Monday)", "Follow-up Q&A session with Horizon Hobby"),
        ("May 1 (Friday) 11:59 PM CT", "Submissions due (hard deadline)"),
        ("May 2 (Saturday) 2:00 AM CT", "Grace period end; late submissions after this are rejected"),
        ("May 2 to 3", "Submission review and automated scoring"),
        ("May 3 (Sunday) EOD", "Three finalist teams notified by email"),
        ("May 7 (Thursday)", "Final presentations to Horizon Hobby, winners announced"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Milestone"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for d, m in timeline:
        row = table.add_row().cells
        row[0].text = d
        row[1].text = m

    # --- Section 3: User roles ---
    doc.add_heading("3. User roles", level=1)
    roles = [
        (
            "Admin",
            "Full platform access. Can manage all users, upload and replace datasets, "
            "upload grading scripts, edit competition content, view all submissions, "
            "override scores, manage announcements, and promote users to Manager. "
            "Initial admin account is created via the database seed.",
        ),
        (
            "Manager",
            "Near-full access. Can do everything Admin can except delete other admins, "
            "modify admin permissions, or change platform configuration. Intended for "
            "co-organizers, TAs, and course staff. Managers are created by an Admin from "
            "the admin users page. Manager accounts are not restricted to @illinois.edu.",
        ),
        (
            "Team Lead",
            "A Gies student who registers the team. Must use an @illinois.edu email. "
            "Can invite a teammate, sign the NDA, download data, submit deliverables, "
            "and view the team's leaderboard position.",
        ),
        (
            "Team Member",
            "A second Gies student who joins a team via an invitation code. Same "
            "capabilities as Team Lead after joining.",
        ),
    ]
    for name, desc in roles:
        p = doc.add_paragraph()
        p.add_run(f"{name}. ").bold = True
        p.add_run(desc)

    # --- Section 4: Student flow ---
    doc.add_heading("4. Student flow (end to end)", level=1)
    doc.add_paragraph(
        "From the moment a Gies student lands on the competition URL to the moment "
        "they submit their final deliverables, here is the path they take:"
    )
    steps = [
        (
            "Discover",
            "Student visits https://16-59-203-133.sslip.io, sees a hero with a LOSI "
            "22S Sprint Car photo, the competition value proposition, an interactive "
            "timeline, a four-step structure walkthrough, cash prize details, and a "
            "registration CTA.",
        ),
        (
            "Register (Team Lead)",
            "Student clicks Register, enters their full name, @illinois.edu email, "
            "password, and desired team name. They must confirm Gies enrollment via a "
            "checkbox. A non-@illinois.edu email is rejected at the form level.",
        ),
        (
            "Verify email",
            "The system emails a 6-digit verification code. The student enters the code "
            "on the verify screen. Codes expire after 15 minutes and can be resent.",
        ),
        (
            "Dashboard (incomplete)",
            "Student lands on the dashboard showing team status as Incomplete (1 of 2). "
            "A form lets them enter their teammate's @illinois.edu email to send an "
            "invitation. The invitation code (6 alphanumeric characters, 72 hour expiry) "
            "is both emailed to the teammate and shown on the dashboard so the lead can "
            "share it directly if the email is delayed.",
        ),
        (
            "Teammate joins",
            "The invited student clicks the email link or visits /join, enters the "
            "invitation code, their own account details, and the eligibility checkbox. "
            "Their email must match the invitation email. The system enforces one team "
            "per student at the database level.",
        ),
        (
            "Dashboard (complete, NDA pending)",
            "Both members now see team status Complete (2 of 2). A card prompts each "
            "member to review and sign the NDA.",
        ),
        (
            "NDA signing",
            "Each team member individually visits /nda, reviews the NDA PDF in an "
            "embedded viewer, types their full legal name to acknowledge, and checks "
            "an agreement box. The typed name must match their registered name "
            "(case-insensitive, trimmed). The system records the signer, timestamp, IP "
            "address, and a SHA-256 hash of the exact NDA PDF they signed, so if the "
            "NDA is later updated, the old signatures remain valid for the prior version.",
        ),
        (
            "Data download",
            "After both members sign, and once the admin has enabled downloads globally "
            "(toggled on at kickoff), the team can visit /data to download the competition "
            "brief, dataset, and prediction template. Every download is audit-logged.",
        ),
        (
            "Build (offline)",
            "The team builds their predictive model on their own machines. The platform "
            "is hands-off during this phase.",
        ),
        (
            "Submit",
            "At /submit the team sees four upload zones: Prediction file (.xlsx), Code "
            "and model files (.py, .ipynb, .r, or .zip), Methodology (.pdf), and "
            "Presentation (.pptx). Each zone is independent. Teams can replace any "
            "component any number of times before the deadline; full version history is "
            "retained but only the latest version counts for scoring and review. The "
            "prediction file is structurally validated at upload time. On successful "
            "upload, both team members receive a confirmation email.",
        ),
        (
            "Automatic scoring",
            "When a prediction file is uploaded, the system queues a grading job. A "
            "sandboxed FastAPI grader downloads the submission, the answer key, and "
            "(optionally) an admin-uploaded Python grading script; runs the script with "
            "a 60-second timeout and memory limit; and records a weighted MAPE score. "
            "The team receives an email with their score and current leaderboard rank.",
        ),
        (
            "Leaderboard",
            "At /leaderboard the team sees every scored team ordered by weighted MAPE "
            "(lower is better), with LATE and FINALIST badges where applicable. Admins "
            "can freeze the leaderboard to hide scores pending finalist notification.",
        ),
        (
            "Deadline reminders",
            "Teams without a complete submission receive automated email reminders at "
            "48, 6, and 1 hour before the primary deadline.",
        ),
        (
            "Late handling",
            "Submissions between 11:59 PM CT May 1 and 2:00 AM CT May 2 are accepted "
            "but flagged LATE on the leaderboard and in admin views. Submissions after "
            "the grace window are rejected entirely.",
        ),
        (
            "Finalist notification",
            "On May 3, admins mark the top three teams as finalists. Those teams get a "
            "congratulations email with presentation details. Non-finalist teams get a "
            "thank-you email with their final score and rank.",
        ),
        (
            "Finals day",
            "On May 7, finalist teams present to Horizon Hobby judges. Winners are "
            "announced on the leaderboard.",
        ),
    ]
    for title, body in steps:
        p = doc.add_paragraph()
        p.add_run(f"{title}. ").bold = True
        p.add_run(body)

    # --- Section 5: Admin flow ---
    doc.add_heading("5. Admin flow and feature surface", level=1)
    doc.add_paragraph(
        "Admins sign in at the same /login page as students but with pre-created "
        "credentials. After login they are redirected to /admin, a sidebar-driven "
        "control panel with these sections:"
    )
    admin_sections = [
        (
            "Dashboard",
            "Live platform statistics: registered students, teams, complete teams, "
            "NDA signatures, submissions received, scored submissions. Updates on "
            "every page load.",
        ),
        (
            "Users and Teams",
            "Full list of all teams with status badges, members, and creation dates. "
            "Separate section for staff accounts (admins and managers) with create, "
            "demote, and deactivate actions. Form to create new Manager or Admin "
            "accounts with any email domain. Promote-student-to-Manager action for "
            "students without a team. Only Admins (not Managers) can create or promote "
            "staff.",
        ),
        (
            "Announcements",
            "Create, view, and delete announcements. Each announcement has a title and "
            "Markdown body, is visible on the public /announcements page, and is shown "
            "on all team dashboards. Author and timestamp are recorded.",
        ),
        (
            "FAQ",
            "Add, edit, and delete FAQ items. Items are ordered by display order and "
            "appear on the public /faq page. Answers support Markdown.",
        ),
        (
            "Data and Files",
            "Three sub-sections:\n"
            "  NDA: upload or replace the NDA PDF. The SHA-256 hash is recorded per "
            "version so student signatures pin to the exact file they signed.\n"
            "  Data files: upload the competition brief, dataset, and prediction "
            "template. Toggle individual files active or inactive. Toggle the master "
            "'data downloads enabled' switch that gates student access site-wide.\n"
            "  Grader: upload the answer key (confidential) and optionally upload a "
            "custom Python grading script. If no script is uploaded, a built-in default "
            "weighted MAPE grader runs.",
        ),
        (
            "Submissions",
            "Per-team grid showing all four components with filenames, sizes, LATE "
            "badges, and individual download links. Finalist toggle button per team. "
            "CSV manifest export of all latest submissions with presigned download URLs.",
        ),
        (
            "Leaderboard",
            "Visibility toggle (Visible, Frozen, Hidden), Re-score All button, per-team "
            "score override with reason tracking, CSV export, and a recent grading jobs "
            "table showing status, errors, and timing for the last 20 jobs.",
        ),
        (
            "Broadcast",
            "Compose a subject line and Markdown body, pick a recipient scope (All "
            "registered, All complete teams, Finalists only), and send. Each broadcast "
            "is tracked in history with per-team delivery status. Optionally archives "
            "the broadcast as an announcement.",
        ),
        (
            "Settings",
            "Edit the primary submission deadline, grace period end, and registration "
            "close timestamps. Changes take effect immediately and are audit-logged.",
        ),
        (
            "Audit log",
            "Last 500 administrative actions with who, what, when, entity, and JSON "
            "details. Covers every content upload, user modification, score override, "
            "and configuration change.",
        ),
    ]
    for name, desc in admin_sections:
        p = doc.add_paragraph()
        p.add_run(f"{name}. ").bold = True
        p.add_run(desc)

    # --- Section 6: Security ---
    doc.add_heading("6. Security and data handling", level=1)
    sec_bullets = [
        "All traffic served over HTTPS via a valid Let's Encrypt certificate.",
        "Passwords hashed with bcrypt at cost factor 12.",
        "Session cookies are HttpOnly, Secure, SameSite=Lax, with 24-hour expiry. "
        "Sessions are backed by a database table so they can be revoked.",
        "Failed logins are rate-limited and lock the account for five minutes after "
        "five consecutive failures.",
        "Email and team name uniqueness enforced at the database level (case-insensitive).",
        "One-team-per-student is enforced by a unique database constraint.",
        "File uploads validated by extension and size per component type.",
        "The grading script runs in an isolated Docker container as a non-root user "
        "with no inbound network access and a 60-second subprocess timeout.",
        "S3 bucket has versioning enabled, server-side encryption (AES-256), and all "
        "public access blocked. Student downloads use short-TTL presigned URLs issued "
        "only after NDA verification.",
        "The answer key is stored in the same bucket but its file type makes it "
        "invisible to student-facing download routes.",
        "Every administrative action is recorded to an audit log, including the IP "
        "address for NDA signatures.",
        "The EC2 instance has termination protection enabled and requires IMDSv2 for "
        "metadata access. The instance profile scopes S3 access to the one bucket.",
    ]
    for b in sec_bullets:
        doc.add_paragraph(b, style="List Bullet")

    # --- Section 7: Automated email ---
    doc.add_heading("7. Automated email notifications", level=1)
    doc.add_paragraph(
        "The platform sends transactional email via the Gmail API using an OAuth "
        "refresh token. The following triggers are wired up:"
    )
    emails = [
        "Welcome and email verification code (on registration)",
        "Team invitation code and join link (sent to invitee)",
        "Team complete confirmation (sent to both members)",
        "NDA confirmed, data download now available (sent to both members after second sign)",
        "Submission component uploaded confirmation (per upload)",
        "All four components received, submission complete",
        "Score recorded, current leaderboard rank",
        "Scoring error (if grading script fails on a submission)",
        "Deadline reminder at 48 hours, 6 hours, and 1 hour before the deadline",
        "Late submission accepted notice (if uploaded in the grace window)",
        "Finalist congratulations with presentation details",
        "Non-finalist thank-you with final rank and score",
        "Password reset code (if triggered)",
        "Broadcast messages (any admin-composed message to a chosen scope)",
    ]
    for e in emails:
        doc.add_paragraph(e, style="List Bullet")

    # --- Section 8: What is not yet wired ---
    doc.add_heading("8. Operational notes", level=1)
    doc.add_paragraph(
        "The following items are designed and built into the platform but require "
        "one-time action before the competition launches:"
    )
    notes = [
        "Real NDA PDF upload (admin uploads via /admin/content/nda when available).",
        "Real dataset, prediction template, and competition brief upload (admin "
        "uploads via /admin/content/data).",
        "Answer key and optional grading script upload (admin uploads via "
        "/admin/content/grader).",
        "Host-level cron entries to periodically hit the grading job queue and "
        "deadline reminder sweeper endpoints. Each endpoint is protected by a shared "
        "secret header.",
        "Optional: swap the sslip.io-based HTTPS hostname for a real registered "
        "domain. Platform reconfigures in minutes when a domain is provided.",
    ]
    for n in notes:
        doc.add_paragraph(n, style="List Bullet")

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "Generated for the Gies College of Business Supply Chain Analytics "
        "Competition 2026 leadership team."
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(str(out_path))
    print(f"Wrote {out_path} ({out_path.stat().st_size} bytes)")


# ---- Gmail OAuth ----

def get_access_token() -> str:
    if not (CLIENT_ID and CLIENT_SECRET and REFRESH_TOKEN):
        raise RuntimeError(
            "Missing GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET / GMAIL_REFRESH_TOKEN in .env"
        )
    data = urllib.parse.urlencode({
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "refresh_token": REFRESH_TOKEN,
        "grant_type": "refresh_token",
    }).encode("utf-8")
    req = urllib.request.Request(
        "https://oauth2.googleapis.com/token",
        data=data,
        headers={"Content-Type": "application/x-www-form-urlencoded"},
    )
    with urllib.request.urlopen(req) as resp:
        payload = json.loads(resp.read().decode("utf-8"))
    return payload["access_token"]


# ---- Send ----

def send_email_with_attachment(
    access_token: str,
    to: str,
    subject: str,
    body_text: str,
    attachment_path: Path,
) -> None:
    msg = EmailMessage()
    msg["From"] = f"{SENDER_NAME} <{SENDER}>"
    msg["To"] = to
    msg["Subject"] = subject
    msg.set_content(body_text)

    ctype, _ = mimetypes.guess_type(attachment_path.name)
    if ctype is None:
        ctype = "application/octet-stream"
    maintype, subtype = ctype.split("/", 1)
    with open(attachment_path, "rb") as fh:
        msg.add_attachment(
            fh.read(),
            maintype=maintype,
            subtype=subtype,
            filename=attachment_path.name,
        )

    raw = base64.urlsafe_b64encode(bytes(msg)).decode("ascii")
    body = json.dumps({"raw": raw}).encode("utf-8")
    req = urllib.request.Request(
        "https://gmail.googleapis.com/gmail/v1/users/me/messages/send",
        data=body,
        headers={
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json",
        },
    )
    with urllib.request.urlopen(req) as resp:
        result = json.loads(resp.read().decode("utf-8"))
    print(f"Sent: message id {result.get('id')}")


# ---- main ----

def main() -> None:
    out_dir = Path(__file__).resolve().parent.parent / "tmp"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "SCM-Analytics-Competition-2026-Platform-Overview.docx"
    build_doc(out_path)

    token = get_access_token()
    send_email_with_attachment(
        token,
        to=RECIPIENT,
        subject="Platform overview: SCM Analytics Competition 2026",
        body_text=(
            "Hi,\n\n"
            "Attached is a Word document describing the flow and features of the "
            "Supply Chain Analytics Competition 2026 platform. Share with your "
            "colleagues.\n\n"
            "Live URL: https://16-59-203-133.sslip.io\n\n"
            "— Platform overview generator"
        ),
        attachment_path=out_path,
    )


if __name__ == "__main__":
    main()
